import re
import subprocess
import sys
from pathlib import Path


def preprocess_markdown(md_content: str) -> tuple:
    """
    Przetwarza markdown: numeruje wzory i obrazki, zamienia odnośniki.
    Zwraca: (processed_content, equation_map)
    """
    equation_counter = 0
    figure_counter = 0
    equation_map = {}
    
    # Zamień bloki equation na $$ i usuń \label
    def replace_equation(match):
        nonlocal equation_counter
        equation_counter += 1
        
        content = match.group(1)
        
        # Wyciągnij i zapisz label
        label_match = re.search(r'\\label\{(.*?)\}', content)
        if label_match:
            label = label_match.group(1)
            equation_map[label] = equation_counter
            content = re.sub(r'\s*\\label\{.*?\}\s*', '', content)
        
        return f'\n$$\n{content.strip()}\n$$\n\n'
    
    content = re.sub(
        r'\\begin\{equation\}(.*?)\\end\{equation\}',
        replace_equation,
        md_content,
        flags=re.DOTALL
    )
    
    # Zamień \eqref{label} na (numer)
    for label, num in equation_map.items():
        content = re.sub(
            rf'\\eqref\{{{re.escape(label)}\}}',
            f'({num})',
            content
        )
    
    # Dodaj "Fig. N." do podpisów obrazków
    def replace_figure(match):
        nonlocal figure_counter
        figure_counter += 1
        caption = match.group(1).replace('\n', ' ').strip()
        path = match.group(2).strip()
        return f'![Fig. {figure_counter}. {caption}]({path})'
    
    content = re.sub(
        r'!\[(.*?)\]\((.*?)\)',
        replace_figure,
        content,
        flags=re.DOTALL
    )
    
    return content, equation_map


def add_equation_numbers_to_docx(docx_path: str):
    """Dodaje numery wzorów do DOCX po konwersji."""
    try:
        from docx import Document
        from docx.shared import Pt
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        
        doc = Document(docx_path)
        eq_counter = 0
        
        i = 0
        while i < len(doc.paragraphs):
            para = doc.paragraphs[i]
            
            # Sprawdź czy akapit zawiera wzór display (bez tekstu)
            if para._element.xpath('.//m:oMath') and len(para.text.strip()) == 0:
                eq_counter += 1
                
                # Wycentruj wzór
                para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                
                # Dodaj nowy akapit z numerem po prawej
                if i + 1 < len(doc.paragraphs):
                    new_para = doc.paragraphs[i+1].insert_paragraph_before()
                else:
                    new_para = doc.add_paragraph()
                
                new_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                run = new_para.add_run(f'({eq_counter})')
                run.font.size = Pt(11)
                
                i += 2  # Pomiń nowo dodany akapit
            else:
                i += 1
        
        doc.save(docx_path)
        print(f"✓ Dodano {eq_counter} numerów wzorów")
        
    except ImportError:
        print("⚠ Brak python-docx. Zainstaluj: pip install python-docx")
    except Exception as e:
        print(f"⚠ Błąd postprocessingu: {e}")


def create_reference_docx(output_path: str = 'reference.docx'):
    """Tworzy szablon DOCX z wyjustowanym tekstem."""
    try:
        from docx import Document
        from docx.shared import Pt
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        
        doc = Document()
        
        style = doc.styles['Normal']
        style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        style.font.name = 'Times New Roman'
        style.font.size = Pt(12)
        
        doc.save(output_path)
        print(f"✓ Utworzono szablon: {output_path}")
        return output_path
        
    except ImportError:
        print("⚠ Brak python-docx. Zainstaluj: pip install python-docx")
        return None


def convert_md_to_docx(input_md: str, output_docx: str = None, 
                       math_method: str = 'omml',
                       reference_doc: str = None,
                       create_template: bool = True):
    """Konwertuje Markdown na DOCX z numeracją wzorów i obrazków."""
    input_path = Path(input_md)
    
    if not input_path.exists():
        raise FileNotFoundError(f"Plik {input_md} nie istnieje")
    
    # Wczytaj i preprocessuj
    with open(input_path, 'r', encoding='utf-8') as f:
        content = f.read()
    
    processed_content, eq_map = preprocess_markdown(content)
    
    print(f"Znaleziono {len(eq_map)} wzorów")
    
    # Zapisz do tymczasowego pliku
    temp_file = input_path.with_suffix('.tmp.md')
    with open(temp_file, 'w', encoding='utf-8') as f:
        f.write(processed_content)
    
    if output_docx is None:
        output_docx = input_path.with_suffix('.docx')
    
    output_path = Path(output_docx)
    
    # Utwórz szablon z wyjustowaniem jeśli nie podano
    if reference_doc is None and create_template:
        reference_doc = create_reference_docx('reference.docx')
    
    # Komenda Pandoc
    cmd = [
        'pandoc',
        str(temp_file),
        '-o', str(output_path),
        '--from', 'markdown+tex_math_dollars',
        '--to', 'docx',
        '--standalone',
        '--citeproc',
        '--resource-path', f"{input_path.parent}:.",
        '--highlight-style', 'tango',
        '--dpi', '300',
    ]
    
    if math_method == 'webtex':
        cmd.extend(['--webtex=https://latex.codecogs.com/png.latex?'])
    
    if reference_doc and Path(reference_doc).exists():
        cmd.extend(['--reference-doc', reference_doc])
    
    try:
        print(f"Konwertuję {input_path.name} -> {output_path.name}")
        print(f"Metoda wzorów: {math_method}\n")
        
        result = subprocess.run(cmd, capture_output=True, text=True, check=True)
        
        if result.stderr:
            print(f"Pandoc: {result.stderr}")
        
        # Dodaj numery wzorów do DOCX
        add_equation_numbers_to_docx(str(output_path))
        
        # Usuń plik tymczasowy
        temp_file.unlink()
        
        return output_path
        
    except Exception as e:
        if temp_file.exists():
            temp_file.unlink()
        raise


if __name__ == "__main__":
    if len(sys.argv) < 2:
        print("Użycie: python md2docx.py <plik.md> [output.docx] [opcje]")
        print("\nOpcje:")
        print("  --math omml|webtex    Metoda renderowania wzorów (domyślnie: omml)")
        print("  --template plik.docx  Użyj własnego szablonu Word")
        print("  --no-template         Nie twórz automatycznego szablonu")
        print("\nPrzykłady:")
        print("  python md2docx.py article.md")
        print("  python md2docx.py article.md output.docx --math webtex")
        sys.exit(1)
    
    input_file = sys.argv[1]
    
    if not Path(input_file).exists():
        print(f"✗ Błąd: Plik {input_file} nie istnieje!")
        sys.exit(1)
    
    # Parsuj argumenty
    output_file = None
    math_method = 'omml'
    reference_doc = None
    create_template = True
    
    i = 2
    while i < len(sys.argv):
        arg = sys.argv[i]
        
        if arg == '--math' and i + 1 < len(sys.argv):
            math_method = sys.argv[i + 1]
            i += 2
        elif arg == '--template' and i + 1 < len(sys.argv):
            reference_doc = sys.argv[i + 1]
            i += 2
        elif arg == '--no-template':
            create_template = False
            i += 1
        elif not arg.startswith('--'):
            output_file = arg
            i += 1
        else:
            i += 1
    
    # Konwertuj
    try:
        result = convert_md_to_docx(
            input_file, 
            output_file, 
            math_method=math_method,
            reference_doc=reference_doc,
            create_template=create_template
        )
        print(f"\n✓ Gotowe! Plik: {result}")
    except Exception as e:
        print(f"\n✗ Błąd: {e}")
        sys.exit(1)