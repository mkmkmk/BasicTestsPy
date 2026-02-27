"""
Konwerter bibliografii MS Word XML do DOCX

Skrypt konwertuje plik XML z bibliografią MS Word (format Office Open XML) pobrane np. z Mendeley
na dokument DOCX z sformatowaną listą publikacji.

Użycie:
    python mendeley2docx.py plik.xml

Argumenty:
    plik.xml - ścieżka do pliku XML z bibliografią MS Word

Wyjście:
    Tworzy plik DOCX o tej samej nazwie co plik wejściowy
    (np. publikacje.xml -> publikacje.docx)

Wymagania:
    - python-docx
    - lxml (opcjonalnie, dla lepszego parsowania XML)

Instalacja zależności:
    pip install python-docx

Autor: MK
Data: 2025-02-27
"""

import sys
from docx import Document
from docx.shared import Pt, Inches
import xml.etree.ElementTree as ET

def main(input_file):

    output_file = input_file.replace('.xml', '.docx')

    tree = ET.parse(input_file)
    root = tree.getroot()

    # Namespace dla MS Office
    ns = {'b': 'http://schemas.openxmlformats.org/officeDocument/2006/bibliography'}

    # Utwórz dokument Word
    doc = Document()
    doc.add_heading('Bibliografia', 0)

    # Parsuj źródła
    for source in root.findall('.//b:Source', ns):
        # Pobierz dane
        title = source.find('.//b:Title', ns)
        year = source.find('.//b:Year', ns)
        journal = source.find('.//b:JournalName', ns)
        volume = source.find('.//b:Volume', ns)
        issue = source.find('.//b:Issue', ns)
        pages = source.find('.//b:Pages', ns)

        # Pobierz autorów
        authors = []
        author_list = source.find('.//b:Author', ns)
        if author_list is not None:
            for person in author_list.findall('.//b:Person', ns):
                last = person.find('.//b:Last', ns)
                first = person.find('.//b:First', ns)
                if last is not None and first is not None:
                    authors.append(f"{last.text}, {first.text}")

        # Formatuj wpis bibliograficzny
        entry = ""
        if authors:
            entry += ", ".join(authors) + ". "
        if year is not None:
            entry += f"({year.text}). "
        if title is not None:
            entry += f"{title.text}. "
        if journal is not None:
            entry += f"*{journal.text}*, "
        if volume is not None:
            entry += f"{volume.text}"
        if issue is not None:
            entry += f"({issue.text})"
        if pages is not None:
            entry += f", {pages.text}"

        # Dodaj do dokumentu
        p = doc.add_paragraph(entry)
        p.style = 'List Bullet'

    # Zapisz dokument
    doc.save(output_file)
    print(f"Dokument został utworzony: {output_file}")


if __name__ == "__main__":
    if len(sys.argv) < 2:
        print("Użycie: python mendeley2docx.py plik.xml")
        sys.exit(1)

    input_file = sys.argv[1]
    main(input_file)
